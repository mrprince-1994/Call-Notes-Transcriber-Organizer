import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import NOTES_BASE_DIR


def _md_to_docx(customer_name: str, notes_content: str, filepath: str):
    """Convert markdown notes to a formatted .docx file."""
    doc = Document()

    # Title
    title = doc.add_heading(f"Call Notes — {customer_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
    date_run.font.size = Pt(11)
    date_run.italic = True
    doc.add_paragraph()  # spacer

    for line in notes_content.split("\n"):
        stripped = line.strip()

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif re.match(r"^\d+\.\s", stripped):
            # Numbered list item
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, stripped[2:])
        elif stripped:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)

    doc.save(filepath)


def _add_formatted_text(paragraph, text):
    """Handle bold (**text**) and italic (*text*) markdown in a paragraph."""
    # Split on bold markers first
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic within non-bold parts
            sub_parts = re.split(r"(\*.*?\*)", part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*"):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub)


def save_notes(customer_name: str, notes_content: str) -> str:
    """
    Save notes as .docx to: Call Notes/{customer_name}/{customer_name}_notes_{N}_{datetime}.docx
    Returns the full path of the saved file.
    """
    customer_dir = os.path.join(NOTES_BASE_DIR, customer_name)
    os.makedirs(customer_dir, exist_ok=True)

    existing = [f for f in os.listdir(customer_dir) if f.endswith(".docx")]
    next_num = len(existing) + 1

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    filename = f"{customer_name}_notes_{next_num}_{timestamp}.docx"
    filepath = os.path.join(customer_dir, filename)

    _md_to_docx(customer_name, notes_content, filepath)

    return filepath
