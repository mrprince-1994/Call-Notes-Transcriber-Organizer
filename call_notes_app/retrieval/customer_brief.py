"""Customer Brief Generator — researches a company and produces a formatted DOCX brief.

Follows the AWS Customer Meeting Preparation steering guide:
1. Company Research (revenue, headcount, market cap, growth, business model, service lines)
2. Leadership Research (executives, backgrounds, areas of focus)
3. Technology Landscape (digital transformation, AI/ML, cloud strategy, tech hires)
4. AI/ML Use Cases (earnings calls, press releases, job postings, conferences)
5. AWS Customer References (tiered: Tier 1 highly relevant, Tier 2 adjacent)
6. AWS Solutions Mapping (customer priorities → AWS services)
7. Competitive Context (landscape, technology differentiation)
"""
import json
import os
import re
from datetime import datetime
import boto3
from botocore.config import Config
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import AWS_REGION, NOTES_BASE_DIR

MODEL_ID = "us.anthropic.claude-sonnet-4-6"
NAVY = RGBColor(0x1E, 0x27, 0x61)
RED_CONF = RGBColor(0x99, 0x00, 0x11)

RESEARCH_PROMPT = """You are an expert business analyst preparing a customer meeting brief for an AWS account team.
Follow this research sequence thoroughly:

1. COMPANY RESEARCH — Extract revenue, headcount, market cap, growth rate, business model, service lines, recent earnings, press releases, SEC filings.
2. LEADERSHIP RESEARCH — Identify key executives (CEO, CTO, CIO, CDO, VP Eng). Find backgrounds, previous companies, areas of focus.
3. TECHNOLOGY LANDSCAPE — Digital transformation initiatives, cloud strategy, AI/ML investments, recent technology hires (CTO/CIO/CDO appointments).
4. AI/ML USE CASES — Existing implementations disclosed in earnings calls, press releases, job postings, or conference presentations.
5. AWS CUSTOMER REFERENCES — Case studies from the customer's industry. Tier 1: highly relevant to their priorities. Tier 2: adjacent/related.
6. AWS SOLUTIONS MAPPING — Map discussion topics to specific AWS services and features. Be specific.
7. COMPETITIVE CONTEXT — Competitive landscape and how technology/AI differentiates them.

Return ONLY valid JSON with this exact structure:
{
  "company_name": "Full legal company name",
  "domain": "company.com",
  "overview": "2-3 paragraph company overview",
  "key_facts": {
    "founded": "Year",
    "headquarters": "City, State",
    "employees": "Headcount",
    "revenue": "Latest annual revenue",
    "market_cap": "If public",
    "industry": "Primary industry",
    "business_model": "B2B/B2C/etc",
    "ticker": "Stock ticker if public, or Private"
  },
  "service_lines": ["Main products/services"],
  "financial_snapshot": {
    "revenue": "Latest annual revenue with year",
    "revenue_growth": "YoY growth rate",
    "profitability": "Net income or EBITDA if available",
    "key_metrics": "Any notable KPIs (ARR, DAU, etc.)",
    "outlook": "Forward guidance or analyst consensus"
  },
  "leadership": [
    {"name": "Name", "title": "Title", "background": "Bio with previous companies and areas of focus"}
  ],
  "technology_landscape": "2-3 paragraphs on tech stack, digital transformation, cloud strategy, AI/ML investments, recent tech hires",
  "ai_ml_use_cases": [
    {"name": "Use case name", "description": "What it does", "category": "NLP/CV/Predictive/GenAI/etc"}
  ],
  "aws_customer_references": {
    "tier_1": [
      {"company": "Company name", "what_they_built": "Description", "results": "Metrics/outcomes", "relevance": "Why relevant to this customer"}
    ],
    "tier_2": [
      {"company": "Company name", "what_they_built": "Description", "results": "Metrics/outcomes", "relevance": "Why relevant"}
    ]
  },
  "aws_solutions_alignment": [
    {"customer_priority": "What they need", "aws_service": "Specific AWS service", "value_proposition": "How it helps"}
  ],
  "competitive_context": "1-2 paragraphs on competitive landscape and tech differentiation",
  "discovery_questions": {
    "ai_strategy": ["4-5 open-ended questions about AI maturity, vision, prioritization"],
    "workflow_operations": ["4-5 questions about current workflows, pain points, automation"],
    "data_governance": ["4-5 questions about data strategy, governance, build-vs-buy"],
    "partnership": ["4-5 questions about partnership criteria, ROI measurement, success metrics"],
    "vision": ["4-5 questions about 2-3 year horizon, strategic bets, transformation goals"]
  },
  "meeting_agenda": [
    {"topic": "Topic", "duration": "X min", "details": "What to cover"}
  ],
  "sources": ["URLs or source descriptions"]
}

DISCOVERY QUESTION PRINCIPLES:
- All questions must be open-ended, never yes/no
- Include questions about AI maturity, prioritization criteria, ROI measurement, build-vs-buy philosophy
- Always include questions about the customer's 2-3 year vision
- Surface pain points without being presumptuous

MEETING AGENDA PRINCIPLES:
- Start with customer sharing their vision (let them lead)
- AWS capabilities come after understanding their priorities
- Include 1-2 customer reference deep-dives with concrete metrics
- Include collaborative brainstorm time
- End with concrete next steps and action items
- Total duration: 60-90 minutes"""


def _research_company(company_name: str, domain: str, on_status=None) -> dict:
    """Call Claude to research the company and return structured JSON."""
    if on_status:
        on_status("Researching company...")

    client = boto3.client(
        "bedrock-runtime", region_name=AWS_REGION,
        config=Config(read_timeout=300),
    )

    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 16000,
        "system": RESEARCH_PROMPT,
        "messages": [{
            "role": "user",
            "content": f"Research this company and produce the JSON brief:\n\nCompany: {company_name}\nDomain: {domain}",
        }],
    }

    response = client.invoke_model(
        modelId=MODEL_ID,
        contentType="application/json",
        accept="application/json",
        body=json.dumps(payload),
    )

    result = json.loads(response["body"].read())
    text = result["content"][0]["text"]

    # Extract JSON from response
    json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if json_match:
        text = json_match.group(1)
    start = text.find('{')
    end = text.rfind('}') + 1
    if start >= 0 and end > start:
        text = text[start:end]

    return json.loads(text)


def _add_table(doc, headers, rows):
    """Helper to add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def _build_docx(data: dict, on_status=None) -> str:
    """Build a formatted DOCX brief from the research data. Returns file path."""
    if on_status:
        on_status("Building document...")

    company = data.get("company_name", "Company")
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = NAVY

    # ── 1. Title Page ──
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company)
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Customer Meeting Brief")
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — For Internal AWS Use Only")
    run.font.color.rgb = RED_CONF
    run.font.size = Pt(10)
    run.bold = True

    doc.add_page_break()

    # ── 2. Table of Contents placeholder ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Company Overview", "2. Financial Snapshot", "3. Leadership Team",
        "4. Technology & AI/ML Landscape", "5. Existing AI/ML Use Cases",
        "6. AWS Customer References", "7. AWS Solutions Alignment",
        "8. Discovery Questions", "9. Meeting Agenda Recommendations",
        "10. Appendix: Sources & References",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ── 3. Company Overview ──
    doc.add_heading("1. Company Overview", level=1)
    doc.add_paragraph(data.get("overview", ""))

    # Key Facts table
    doc.add_heading("Key Facts", level=2)
    facts = data.get("key_facts", {})
    if facts:
        rows = [[k.replace("_", " ").title(), str(v)] for k, v in facts.items()]
        t = _add_table(doc, ["Attribute", "Value"], rows)
        # Bold the label column
        for row in t.rows[1:]:
            for r in row.cells[0].paragraphs[0].runs:
                r.bold = True

    services = data.get("service_lines", [])
    if services:
        doc.add_heading("Service Lines", level=2)
        for s in services:
            doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # ── 4. Financial Snapshot ──
    fin = data.get("financial_snapshot", {})
    if fin:
        doc.add_heading("2. Financial Snapshot", level=1)
        rows = [[k.replace("_", " ").title(), str(v)] for k, v in fin.items()]
        _add_table(doc, ["Metric", "Value"], rows)
        doc.add_page_break()

    # ── 5. Leadership Team ──
    leaders = data.get("leadership", [])
    if leaders:
        doc.add_heading("3. Leadership Team", level=1)
        for leader in leaders:
            p = doc.add_paragraph()
            run = p.add_run(f"{leader.get('name', '')} — {leader.get('title', '')}")
            run.bold = True
            run.font.size = Pt(11)
            doc.add_paragraph(leader.get("background", ""))
        doc.add_page_break()

    # ── 6. Technology & AI/ML Landscape ──
    doc.add_heading("4. Technology & AI/ML Landscape", level=1)
    doc.add_paragraph(data.get("technology_landscape", ""))

    # ── 7. Existing AI/ML Use Cases ──
    use_cases = data.get("ai_ml_use_cases", [])
    if use_cases:
        doc.add_heading("5. Existing AI/ML Use Cases", level=1)
        rows = [[uc.get("name", ""), uc.get("description", ""), uc.get("category", "")]
                for uc in use_cases]
        _add_table(doc, ["Use Case", "Description", "Category"], rows)

    doc.add_page_break()

    # ── 8. AWS Customer References (Tiered) ──
    refs = data.get("aws_customer_references", {})
    tier1 = refs.get("tier_1", [])
    tier2 = refs.get("tier_2", [])
    if tier1 or tier2:
        doc.add_heading("6. AWS Customer References", level=1)
        if tier1:
            doc.add_heading("Tier 1 — Highly Relevant", level=2)
            rows = [[r.get("company", ""), r.get("what_they_built", ""),
                     r.get("results", ""), r.get("relevance", "")] for r in tier1]
            _add_table(doc, ["Company", "What They Built", "Results/Metrics", "Relevance"], rows)
        if tier2:
            doc.add_heading("Tier 2 — Adjacent / Related", level=2)
            rows = [[r.get("company", ""), r.get("what_they_built", ""),
                     r.get("results", ""), r.get("relevance", "")] for r in tier2]
            _add_table(doc, ["Company", "What They Built", "Results/Metrics", "Relevance"], rows)
        doc.add_page_break()

    # ── 9. AWS Solutions Alignment ──
    alignment = data.get("aws_solutions_alignment", [])
    if alignment:
        doc.add_heading("7. AWS Solutions Alignment", level=1)
        rows = [[a.get("customer_priority", ""), a.get("aws_service", ""),
                 a.get("value_proposition", "")] for a in alignment]
        _add_table(doc, ["Customer Priority", "AWS Service", "Value Proposition"], rows)
        doc.add_page_break()

    # ── 10. Competitive Context ──
    doc.add_heading("Competitive Context", level=1)
    doc.add_paragraph(data.get("competitive_context", ""))

    # ── 11. Discovery Questions ──
    questions = data.get("discovery_questions", {})
    if questions:
        doc.add_heading("8. Discovery Questions", level=1)
        for theme, qs in questions.items():
            label = theme.replace("_", " ").title()
            doc.add_heading(label, level=2)
            for q in qs:
                doc.add_paragraph(q, style='List Bullet')
        doc.add_page_break()

    # ── 12. Meeting Agenda ──
    agenda = data.get("meeting_agenda", [])
    if agenda:
        doc.add_heading("9. Meeting Agenda Recommendations", level=1)
        rows = [[a.get("topic", ""), a.get("duration", ""), a.get("details", "")]
                for a in agenda]
        _add_table(doc, ["Topic", "Duration", "Details"], rows)
        doc.add_page_break()

    # ── 13. Sources ──
    sources = data.get("sources", [])
    if sources:
        doc.add_heading("10. Appendix: Sources & References", level=1)
        for s in sources:
            doc.add_paragraph(s, style='List Bullet')

    # ── Save ──
    customer_dir = os.path.join(NOTES_BASE_DIR, company)
    os.makedirs(customer_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M")
    filename = f"{company}_brief_{ts}.docx"
    filepath = os.path.join(customer_dir, filename)
    doc.save(filepath)
    return filepath


def generate_customer_brief(company_name: str, domain: str, on_status=None) -> str:
    """Full pipeline: research -> build DOCX. Returns the saved file path."""
    data = _research_company(company_name, domain, on_status=on_status)
    if on_status:
        on_status("Building document...")
    filepath = _build_docx(data, on_status=on_status)
    return filepath
