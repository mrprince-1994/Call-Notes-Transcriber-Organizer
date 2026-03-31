import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import NOTES_BASE_DIR


def _md_to_docx(customer_name: str, notes_content: str, filepath: str):
    """Convert markdown notes to a formatted .docx file."""
    doc = Document()

    # Title
    title = doc.add_heading(f"Call Notes — {customer_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
    date_run.font.size = Pt(11)
    date_run.italic = True
    doc.add_paragraph()  # spacer

    for line in notes_content.split("\n"):
        stripped = line.strip()

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif re.match(r"^\d+\.\s", stripped):
            # Numbered list item
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, stripped[2:])
        elif stripped:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)

    doc.save(filepath)


def _add_formatted_text(paragraph, text):
    """Handle bold (**text**) and italic (*text*) markdown in a paragraph."""
    # Split on bold markers first
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic within non-bold parts
            sub_parts = re.split(r"(\*.*?\*)", part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*"):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub)


def export_share_html(customer_name: str, notes_content: str, filepath: str):
    """Export notes as a self-contained, styled HTML file for sharing."""
    import html as html_mod
    timestamp = datetime.now().strftime("%B %d, %Y %I:%M %p")

    # Convert markdown to simple HTML
    lines = notes_content.split("\n")
    body_parts = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            body_parts.append("<br>")
        elif stripped.startswith("### "):
            body_parts.append(f"<h3>{html_mod.escape(stripped[4:])}</h3>")
        elif stripped.startswith("## "):
            body_parts.append(f"<h2>{html_mod.escape(stripped[3:])}</h2>")
        elif stripped.startswith("# "):
            body_parts.append(f"<h1>{html_mod.escape(stripped[2:])}</h1>")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            body_parts.append(f"<li>{html_mod.escape(stripped[2:])}</li>")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            body_parts.append(f"<li>{html_mod.escape(text)}</li>")
        else:
            # Handle bold
            escaped = html_mod.escape(stripped)
            escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
            escaped = re.sub(r"\*(.+?)\*", r"<em>\1</em>", escaped)
            body_parts.append(f"<p>{escaped}</p>")

    body_html = "\n".join(body_parts)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Call Notes — {html_mod.escape(customer_name)}</title>
<style>
  body {{ font-family: 'Segoe UI', system-ui, sans-serif; max-width: 800px; margin: 40px auto;
         padding: 0 24px; background: #0a0a0a; color: #c8ccd0; line-height: 1.6; }}
  h1 {{ color: #10a37f; border-bottom: 2px solid #1f2937; padding-bottom: 8px; }}
  h2 {{ color: #eef0f2; margin-top: 24px; }}
  h3 {{ color: #eef0f2; }}
  li {{ margin: 4px 0; }}
  strong {{ color: #eef0f2; }}
  .meta {{ color: #5a6270; font-size: 0.9em; margin-bottom: 24px; }}
  .footer {{ margin-top: 40px; padding-top: 16px; border-top: 1px solid #1f2937;
             color: #5a6270; font-size: 0.8em; }}
  @media (prefers-color-scheme: light) {{
    body {{ background: #ffffff; color: #1a1a1a; }}
    h1 {{ color: #0d8c6d; border-bottom-color: #e5e7eb; }}
    h2, h3, strong {{ color: #111827; }}
    .meta, .footer {{ color: #6b7280; }}
    .footer {{ border-top-color: #e5e7eb; }}
  }}
</style>
</head>
<body>
<h1>Call Notes — {html_mod.escape(customer_name)}</h1>
<div class="meta">{html_mod.escape(timestamp)}</div>
{body_html}
<div class="footer">Generated by Sales Productivity Suite</div>
</body>
</html>"""

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)


def save_notes(customer_name: str, notes_content: str) -> str:
    """
    Save notes as .md to: Call Notes/{customer_name}/{customer_name}_notes_{N}_{datetime}.md
    Returns the full path of the saved file.
    """
    customer_dir = os.path.join(NOTES_BASE_DIR, customer_name)
    os.makedirs(customer_dir, exist_ok=True)

    existing = [f for f in os.listdir(customer_dir) if f.endswith(".md")]
    next_num = len(existing) + 1

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    filename = f"{customer_name}_notes_{next_num}_{timestamp}.md"
    filepath = os.path.join(customer_dir, filename)

    header = f"# Call Notes — {customer_name}\n\n_Date: {datetime.now().strftime('%B %d, %Y %I:%M %p')}_\n\n---\n\n"
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(header + notes_content)

    # Auto-generate .docx alongside the .md
    docx_path = filepath.replace(".md", ".docx")
    try:
        _md_to_docx(customer_name, notes_content, docx_path)
    except Exception as e:
        print(f"[storage] Warning: failed to auto-generate docx: {e}")

    return filepath
